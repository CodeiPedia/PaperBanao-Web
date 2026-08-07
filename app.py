import streamlit as st
import requests
import fitz  # PyMuPDF
import os
import markdown
from datetime import datetime, timedelta, timezone
import re
import uuid
import base64
import bcrypt
import time
import random
import smtplib
import razorpay
import logging
from xhtml2pdf import pisa
from PIL import Image
from email.mime.text import MIMEText
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from supabase import create_client, Client

# --- LOGGING CONFIGURATION ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Page Config ---
st.set_page_config(page_title="PaperBanao - AI Question Paper", page_icon="📝", layout="centered")

# ==========================================
# --- 🛑 SECRETS: PULLING KEYS SECURELY ---
# ==========================================
SERVER_API_KEY = st.secrets["GEMINI_API_KEY"]
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]

# ==========================================
# --- INITIALIZE SUPABASE CLIENT ---
# ==========================================
@st.cache_resource
def init_supabase():
    try:
        return create_client(SUPABASE_URL, SUPABASE_KEY)
    except Exception as e:
        logging.error(f"[Database Connection Error] {e}")
        st.error("Database Connection Error. Please contact support.")
        return None

supabase: Client = init_supabase()

class StoredLogo(BytesIO):
    def __init__(self, data: bytes, mimetype: str):
        super().__init__(data)
        self.type = mimetype

# --- INITIALIZE RAZORPAY CLIENT ---
@st.cache_resource
def init_razorpay():
    try:
        return razorpay.Client(auth=(st.secrets["RAZORPAY_KEY_ID"], st.secrets["RAZORPAY_KEY_SECRET"]))
    except Exception as e:
        logging.error(f"[Razorpay Init Error] {e}")
        return None

razorpay_client = init_razorpay()
APP_URL = "https://paperbanao-web.streamlit.app/"
PRO_PRICE_INR = 99
PRO_DURATION_DAYS = 30

# --- DB HELPER FUNCTIONS ---
def hash_password(password):
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(password, stored_hash):
    if stored_hash.startswith("$2b$") or stored_hash.startswith("$2a$"):
        try:
            return bcrypt.checkpw(password.encode(), stored_hash.encode())
        except ValueError:
            return False
    import hashlib
    legacy_hash = hashlib.sha256(password.encode()).hexdigest()
    return legacy_hash == stored_hash

def create_user(username, password, email):
    username = username.strip()
    email = email.strip().lower()
    try:
        existing = supabase.table("users").select("username").ilike("username", username).execute()
        if existing.data:
            return False, "Username already exists. Choose another."
        existing_email = supabase.table("users").select("username").ilike("email", email).execute()
        if existing_email.data:
            return False, "An account with this email already exists."
        data = {"username": username, "password": hash_password(password), "email": email, "papers_generated": 0, "is_pro": False}
        supabase.table("users").insert(data).execute()
        return True, "Account created successfully! Please Login."
    except Exception as e:
        logging.error(f"[Signup Error] {e}")
        if st.secrets.get("DEBUG_MODE", False):
            return False, f"DEBUG: {e}"
        return False, "Something went wrong creating your account. Please try again."

def authenticate_user(username, password):
    username = username.strip()
    try:
        res = supabase.table("users").select("*").eq("username", username).execute()
    except Exception as e:
        logging.error(f"[Auth Error] {e}")
        st.error("Login is temporarily unavailable. Please try again shortly.")
        return None
    if not res.data:
        return None
    user = res.data[0]
    if not verify_password(password, user["password"]):
        return None
    if not (user["password"].startswith("$2b$") or user["password"].startswith("$2a$")):
        try:
            supabase.table("users").update({"password": hash_password(password)}).eq("username", username).execute()
        except Exception as e:
            logging.error(f"[Password Upgrade Error] {e}")
    return user

def get_user_data(username):
    try:
        res = supabase.table("users").select("papers_generated, is_pro, email, pro_expires_at").eq("username", username).execute()
        if len(res.data) > 0:
            row = res.data[0]
            expires_at = row.get("pro_expires_at")
            effective_pro = False
            if expires_at:
                try:
                    effective_pro = datetime.fromisoformat(expires_at) > datetime.now(timezone.utc)
                except ValueError:
                    effective_pro = False
            return {
                "papers_generated": row["papers_generated"],
                "is_pro": effective_pro,
                "email": row.get("email"),
                "pro_expires_at": expires_at
            }
    except Exception as e:
        logging.error(f"[get_user_data Error] {e}")
    return {"papers_generated": 0, "is_pro": False, "email": None, "pro_expires_at": None}

def update_paper_count(username):
    try:
        current_count = get_user_data(username)["papers_generated"]
        supabase.table("users").update({"papers_generated": current_count + 1}).eq("username", username).execute()
    except Exception as e:
        logging.error(f"[update_paper_count Error] {e}")

def delete_paper(paper_id, username):
    try:
        supabase.table("papers").delete().eq("id", paper_id).eq("username", username).execute()
    except Exception as e:
        logging.error(f"[delete_paper Error] {e}")
        st.error("Couldn't delete that paper. Please try again.")

def get_institution_defaults(username):
    try:
        res = supabase.table("users").select(
            "default_inst_name, default_inst_address, default_inst_contact, "
            "default_teacher_name, default_paper_language, default_board_format, "
            "default_logo_base64, default_logo_mimetype"
        ).eq("username", username).execute()
        if res.data:
            return res.data[0]
    except Exception as e:
        logging.error(f"[get_institution_defaults Error] {e}")
    return {}

def save_institution_defaults(username, inst_name, inst_address, inst_contact, teacher_name,
                               paper_language, board_format, logo_bytes=None, logo_mimetype=None):
    try:
        update_data = {
            "default_inst_name": inst_name,
            "default_inst_address": inst_address,
            "default_inst_contact": inst_contact,
            "default_teacher_name": teacher_name,
            "default_paper_language": paper_language,
            "default_board_format": board_format,
        }
        if logo_bytes is not None:
            update_data["default_logo_base64"] = base64.b64encode(logo_bytes).decode()
            update_data["default_logo_mimetype"] = logo_mimetype
        supabase.table("users").update(update_data).eq("username", username).execute()
        return True
    except Exception as e:
        logging.error(f"[save_institution_defaults Error] {e}")
        return False

# --- CURRICULUM (Class -> Subject -> Chapters), shared across all users ---
# This is deliberately teacher-maintained rather than hardcoded: NCERT/BSEB
# syllabi change over time, so a hardcoded list would silently go stale.
# Anyone can add/extend a class+subject's chapter list; it's saved once and
# reused by everyone after that.
CLASS_OPTIONS = [f"Class {i}" for i in range(1, 13)]

def get_subjects_for_class(class_name):
    try:
        res = supabase.table("curriculum").select("subject_name").eq("class_name", class_name).execute()
        return sorted(set(r["subject_name"] for r in res.data))
    except Exception as e:
        logging.error(f"[get_subjects_for_class Error] {e}")
        return []

def get_chapters(class_name, subject_name):
    try:
        res = supabase.table("curriculum").select("chapters").eq("class_name", class_name).eq("subject_name", subject_name).execute()
        if res.data:
            return [c.strip() for c in res.data[0]["chapters"].split(",") if c.strip()]
    except Exception as e:
        logging.error(f"[get_chapters Error] {e}")
    return []

def save_chapters(class_name, subject_name, chapters_list):
    try:
        chapters_str = ", ".join(sorted(set(c.strip() for c in chapters_list if c.strip())))
        existing = supabase.table("curriculum").select("id").eq("class_name", class_name).eq("subject_name", subject_name).execute()
        if existing.data:
            supabase.table("curriculum").update({"chapters": chapters_str, "updated_at": datetime.now(timezone.utc).isoformat()}).eq("id", existing.data[0]["id"]).execute()
        else:
            supabase.table("curriculum").insert({"class_name": class_name, "subject_name": subject_name, "chapters": chapters_str}).execute()
        return True
    except Exception as e:
        logging.error(f"[save_chapters Error] {e}")
        return False

# --- PASSWORD RESET (Email OTP) ---
def send_otp_email(to_email, otp):
    try:
        smtp_host = st.secrets.get("SMTP_HOST", "smtp.gmail.com")
        smtp_port = int(st.secrets.get("SMTP_PORT", 465))
        smtp_user = st.secrets["SMTP_USER"]
        smtp_pass = st.secrets["SMTP_PASSWORD"]

        msg = MIMEText(
            f"Your PaperBanao password reset code is: {otp}\n\n"
            f"This code expires in 10 minutes. If you didn't request this, you can ignore this email."
        )
        msg["Subject"] = "PaperBanao - Password Reset Code"
        msg["From"] = smtp_user
        msg["To"] = to_email

        with smtplib.SMTP_SSL(smtp_host, smtp_port) as server:
            server.login(smtp_user, smtp_pass)
            server.sendmail(smtp_user, [to_email], msg.as_string())
        return True
    except Exception as e:
        logging.error(f"[Email Send Error] {e}")
        return False

def request_password_reset(identifier):
    identifier = identifier.strip()
    generic_msg = "If that account exists, a reset code has been sent to its registered email."
    try:
        res = supabase.table("users").select("username, email") \
            .or_(f"username.eq.{identifier},email.eq.{identifier.lower()}").execute()
    except Exception as e:
        logging.error(f"[Reset Lookup Error] {e}")
        return False, "Something went wrong. Please try again."

    if not res.data:
        return True, generic_msg 

    user = res.data[0]
    if not user.get("email"):
        return False, "This account has no email on file. Please contact support to reset it."

    otp = f"{random.randint(0, 999999):06d}"
    otp_hash = hash_password(otp)  
    expires_at = (datetime.now(timezone.utc) + timedelta(minutes=10)).isoformat()

    try:
        supabase.table("users").update({
            "reset_otp": otp_hash,
            "reset_otp_expires": expires_at
        }).eq("username", user["username"]).execute()
    except Exception as e:
        logging.error(f"[Reset Store Error] {e}")
        return False, "Something went wrong. Please try again."

    if send_otp_email(user["email"], otp):
        return True, generic_msg
    else:
        return False, "Couldn't send the reset email right now. Please try again shortly."

def verify_and_reset_password(identifier, otp, new_password):
    identifier = identifier.strip()
    try:
        res = supabase.table("users").select("username, reset_otp, reset_otp_expires") \
            .or_(f"username.eq.{identifier},email.eq.{identifier.lower()}").execute()
    except Exception as e:
        logging.error(f"[Reset Verify Error] {e}")
        return False, "Something went wrong. Please try again."

    if not res.data:
        return False, "Incorrect code or account."

    user = res.data[0]
    if not user.get("reset_otp") or not user.get("reset_otp_expires"):
        return False, "No active reset request found. Please request a new code."

    if datetime.fromisoformat(user["reset_otp_expires"]) < datetime.now(timezone.utc):
        return False, "This code has expired. Please request a new one."

    try:
        if not bcrypt.checkpw(otp.encode(), user["reset_otp"].encode()):
            return False, "Incorrect code."
    except ValueError:
        return False, "Incorrect code."

    try:
        supabase.table("users").update({
            "password": hash_password(new_password),
            "reset_otp": None,
            "reset_otp_expires": None
        }).eq("username", user["username"]).execute()
    except Exception as e:
        logging.error(f"[Reset Update Error] {e}")
        return False, "Something went wrong. Please try again."

    return True, "Password updated! Please login with your new password."

# --- RAZORPAY: PAYMENT LINK + VERIFICATION ---
def create_pro_payment_link(username, email):
    if not razorpay_client:
        return None, "Payment system isn't configured right now."
    try:
        link_data = {
            "amount": PRO_PRICE_INR * 100, 
            "currency": "INR",
            "accept_partial": False,
            "description": f"PaperBanao Pro - {PRO_DURATION_DAYS} Days",
            "customer": {"name": username, "email": email} if email else {"name": username},
            "notify": {"email": bool(email)},
            "reminder_enable": True,
            "notes": {"username": username},
            "callback_url": APP_URL,
            "callback_method": "get"
        }
        link = razorpay_client.payment_link.create(link_data)
        return link["short_url"], None
    except Exception as e:
        logging.error(f"[Payment Link Error] {e}")
        return None, "Couldn't create the payment link. Please try again."

def process_payment_callback(params):
    required = ["razorpay_payment_link_id", "razorpay_payment_link_reference_id",
                "razorpay_payment_link_status", "razorpay_payment_id", "razorpay_signature"]
    if not all(k in params for k in required):
        return False, None

    try:
        razorpay_client.utility.verify_payment_link_signature(params)
    except razorpay.errors.SignatureVerificationError as e:
        logging.error(f"[Payment Signature Error] {e}")
        return False, "Payment verification failed. If money was deducted, please contact support."

    if params["razorpay_payment_link_status"] != "paid":
        return False, None

    payment_id = params["razorpay_payment_id"]

    try:
        existing = supabase.table("payments").select("payment_id").eq("payment_id", payment_id).execute()
        if existing.data:
            return True, "Payment already processed."
    except Exception as e:
        logging.error(f"[Payment Idempotency Check Error] {e}")
        return False, "Something went wrong verifying your payment. Please contact support."

    try:
        link_details = razorpay_client.payment_link.fetch(params["razorpay_payment_link_id"])
        username = link_details.get("notes", {}).get("username")
    except Exception as e:
        logging.error(f"[Payment Link Fetch Error] {e}")
        return False, "Couldn't confirm payment details. Please contact support."

    if not username:
        return False, "Couldn't identify the account for this payment. Please contact support."

    try:
        user_res = supabase.table("users").select("pro_expires_at").eq("username", username).execute()
        now = datetime.now(timezone.utc)
        current_expiry = None
        if user_res.data and user_res.data[0].get("pro_expires_at"):
            current_expiry = datetime.fromisoformat(user_res.data[0]["pro_expires_at"])
        start_from = current_expiry if (current_expiry and current_expiry > now) else now
        new_expiry = start_from + timedelta(days=PRO_DURATION_DAYS)

        supabase.table("users").update({
            "is_pro": True,
            "pro_expires_at": new_expiry.isoformat()
        }).eq("username", username).execute()

        supabase.table("payments").insert({
            "payment_id": payment_id,
            "username": username,
            "amount_inr": PRO_PRICE_INR
        }).execute()
    except Exception as e:
        logging.error(f"[Payment Credit Error] {e}")
        return False, "Payment received but activation failed. Please contact support with your payment ID: " + payment_id

    return True, f"Payment successful! Pro is active until {new_expiry.strftime('%d %b %Y')}."

# --- INITIALIZE SESSION STATE ---
if "logged_in" not in st.session_state: st.session_state.logged_in = False
if "username" not in st.session_state: st.session_state.username = ""
if "blocks" not in st.session_state: st.session_state.blocks = [] 
if "file_name" not in st.session_state: st.session_state.file_name = "PaperBanao_Exam"
if "current_subject" not in st.session_state: st.session_state.current_subject = "Unknown Subject"
if "current_class" not in st.session_state: st.session_state.current_class = ""
if "current_marks" not in st.session_state: st.session_state.current_marks = ""
if "login_attempts" not in st.session_state: st.session_state.login_attempts = 0
if "login_locked_until" not in st.session_state: st.session_state.login_locked_until = 0

# ==========================================
# --- 🔐 LOGIN & SIGNUP UI ---
# ==========================================
if not st.session_state.logged_in:
    st.markdown("<h1 style='text-align: center;'>📝 PaperBanao AI (Cloud)</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Generate precise question papers in seconds.</p>", unsafe_allow_html=True)
    
    if not supabase:
        st.error("⚠️ SYSTEM ADMIN: Please configure 'SUPABASE_URL' and 'SUPABASE_KEY' in the code to enable Login.")
        st.stop()
        
    st.markdown("---")
    t_login, t_signup, t_forgot = st.tabs(["Login", "Sign Up (Free Trial)", "Forgot Password"])
    
    with t_login:
        l_user = st.text_input("Username", key="l_user")
        l_pass = st.text_input("Password", type="password", key="l_pass")

        locked_remaining = st.session_state.login_locked_until - time.time()
        if locked_remaining > 0:
            st.error(f"Too many failed attempts. Try again in {int(locked_remaining)}s.")
        elif st.button("Login", use_container_width=True):
            user = authenticate_user(l_user, l_pass)
            if user:
                st.session_state.logged_in = True
                st.session_state.username = user["username"]
                st.session_state.login_attempts = 0
                st.rerun()
            else:
                st.session_state.login_attempts += 1
                if st.session_state.login_attempts >= 5:
                    st.session_state.login_locked_until = time.time() + 60
                    st.session_state.login_attempts = 0
                    st.error("Too many failed attempts. Locked for 60s.")
                else:
                    st.error("Invalid Username or Password")
                
    with t_signup:
        s_user = st.text_input("New Username", key="s_user")
        s_email = st.text_input("Email", key="s_email", help="Needed for password reset")
        s_pass = st.text_input("New Password", type="password", key="s_pass")
        if st.button("Create Account & Get 5 Free Papers", use_container_width=True):
            if len(s_user.strip()) < 3:
                st.error("Username must be at least 3 characters.")
            elif not re.match(r'^[A-Za-z0-9_.]+$', s_user.strip()):
                st.error("Username can only contain letters, numbers, underscores, and dots.")
            elif not re.match(r'^[^@\s]+@[^@\s]+\.[^@\s]+$', s_email.strip()):
                st.error("Please enter a valid email address.")
            elif len(s_pass) < 6:
                st.error("Password must be at least 6 characters.")
            else:
                ok, msg = create_user(s_user, s_pass, s_email)
                if ok: st.success(msg)
                else: st.error(msg)

    with t_forgot:
        if "reset_otp_sent" not in st.session_state: st.session_state.reset_otp_sent = False
        if "reset_identifier" not in st.session_state: st.session_state.reset_identifier = ""
        if "reset_request_locked_until" not in st.session_state: st.session_state.reset_request_locked_until = 0

        if not st.session_state.reset_otp_sent:
            st.write("Enter your username or email — we'll send a 6-digit code to your registered email.")
            f_id = st.text_input("Username or Email", key="f_id")

            req_locked_remaining = st.session_state.reset_request_locked_until - time.time()
            if req_locked_remaining > 0:
                st.info(f"Please wait {int(req_locked_remaining)}s before requesting another code.")
            elif st.button("Send Reset Code", use_container_width=True):
                if f_id.strip() == "":
                    st.error("Please enter your username or email.")
                else:
                    ok, msg = request_password_reset(f_id)
                    st.session_state.reset_request_locked_until = time.time() + 60
                    if ok:
                        st.session_state.reset_otp_sent = True
                        st.session_state.reset_identifier = f_id
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
        else:
            st.info("Code sent (if the account exists). Check your email — it expires in 10 minutes.")
            f_otp = st.text_input("6-digit Code", key="f_otp", max_chars=6)
            f_new_pass = st.text_input("New Password", type="password", key="f_new_pass")
            fc1, fc2 = st.columns(2)
            if fc1.button("Reset Password", use_container_width=True):
                if len(f_new_pass) < 6:
                    st.error("Password must be at least 6 characters.")
                else:
                    ok, msg = verify_and_reset_password(st.session_state.reset_identifier, f_otp, f_new_pass)
                    if ok:
                        st.success(msg)
                        st.session_state.reset_otp_sent = False
                    else:
                        st.error(msg)
            if fc2.button("Start Over", use_container_width=True):
                st.session_state.reset_otp_sent = False
                st.rerun()
    st.stop()

# ==========================================
# --- APP LOGIC (IF LOGGED IN) ---
# ==========================================

qp = st.query_params
if "razorpay_payment_id" in qp:
    ok, msg = process_payment_callback(dict(qp))
    st.query_params.clear()
    if ok:
        st.success(msg)
    elif msg:
        st.error(msg)

user_data = get_user_data(st.session_state.username)
papers_used = user_data["papers_generated"]
is_pro = user_data["is_pro"]
pro_expires_at = user_data["pro_expires_at"]
user_email = user_data["email"]
FREE_LIMIT = 5

col_logo, col_title, col_logout = st.columns([1, 4, 1])
with col_logo: st.markdown("<h1>📝</h1>", unsafe_allow_html=True) 
with col_title: st.title("PaperBanao")
with col_logout:
    st.write(f"👤 **{st.session_state.username}**")
    if st.button("Logout"):
        st.session_state.logged_in = False
        st.session_state.username = ""
        st.session_state.blocks = []
        st.session_state.blocks_saved = True
        st.session_state.confirm_overwrite = False
        if "inst_defaults" in st.session_state: del st.session_state["inst_defaults"]
        st.rerun()

# ==========================================
# --- SIDEBAR & SETTINGS ---
# ==========================================
st.sidebar.header("💳 Your Account")
if is_pro:
    expiry_str = datetime.fromisoformat(pro_expires_at).strftime("%d %b %Y")
    st.sidebar.success(f"🌟 PRO Member (Unlimited)\n\nActive until {expiry_str}")
    if st.sidebar.button("Renew Pro (₹99 / 30 days)", use_container_width=True):
        with st.spinner("Creating payment link..."):
            link, err = create_pro_payment_link(st.session_state.username, user_email)
        if link:
            st.sidebar.link_button("Click to Pay ₹99", link, use_container_width=True)
        else:
            st.sidebar.error(err)
else:
    papers_left = FREE_LIMIT - papers_used
    st.sidebar.info(f"🪙 Free Credits: {papers_left} / {FREE_LIMIT}")
    st.sidebar.progress(papers_used / FREE_LIMIT if papers_used <= FREE_LIMIT else 1.0)
    if papers_left <= 0:
        st.sidebar.error("⚠️ Free Trial Expired!")
    if st.sidebar.button("⬆️ Upgrade to Pro (₹99 / 30 days)", use_container_width=True):
        with st.spinner("Creating payment link..."):
            link, err = create_pro_payment_link(st.session_state.username, user_email)
        if link:
            st.sidebar.link_button("Click to Pay ₹99", link, use_container_width=True)
        else:
            st.sidebar.error(err)

st.sidebar.markdown("---")
st.sidebar.header("⚙️ Advanced Settings")
st.sidebar.write("Use your own free Gemini API Key when the server limit is reached.")
user_api_key = st.sidebar.text_input("Your Gemini API Key (Optional)", type="password", help="Get your free key from Google AI Studio")

if user_api_key:
    st.sidebar.success("✅ Personal API Key Active!")

st.sidebar.markdown("---")
st.sidebar.header("🏫 Institute Details")

if "inst_defaults" not in st.session_state:
    st.session_state.inst_defaults = get_institution_defaults(st.session_state.username)
_d = st.session_state.inst_defaults

inst_logo_upload = st.sidebar.file_uploader("Upload Logo", type=["png", "jpg", "jpeg"], help="Leave empty to keep your saved default logo")
inst_name = st.sidebar.text_input("Institute Name", value=_d.get("default_inst_name") or "My Success Academy")
exam_time = st.sidebar.text_input("Exam Time", value="2 Hours")

st.sidebar.markdown("---")
st.sidebar.header("🏢 Footer Details")
teacher_name = st.sidebar.text_input("Teacher Name", value=_d.get("default_teacher_name") or "Mr. Suraj")
inst_address = st.sidebar.text_input("Institute Address", value=_d.get("default_inst_address") or "NH-22 Education Lane, City")
inst_contact = st.sidebar.text_input("Contact Number", value=_d.get("default_inst_contact") or "+91 9310038172")

st.sidebar.markdown("---")
st.sidebar.header("📜 Formatting")
_board_options = ["Standard", "BSEB (Bihar Board)", "CBSE", "ICSE"]
_lang_options = ["English", "Hindi", "Bilingual"]
board_format = st.sidebar.selectbox("Board Pattern", _board_options,
    index=_board_options.index(_d["default_board_format"]) if _d.get("default_board_format") in _board_options else 0)
paper_language = st.sidebar.selectbox("Paper Language", _lang_options,
    index=_lang_options.index(_d["default_paper_language"]) if _d.get("default_paper_language") in _lang_options else 0)
include_answer_key = st.sidebar.toggle("Include Answer Key", value=True)
is_two_column = st.sidebar.toggle("📄 Two-Column Format", value=True)

if st.sidebar.button("💾 Save these as my Default", use_container_width=True):
    logo_bytes, logo_mimetype = None, None
    if inst_logo_upload is not None:
        inst_logo_upload.seek(0)
        logo_bytes = inst_logo_upload.getvalue()
        logo_mimetype = inst_logo_upload.type
    if save_institution_defaults(st.session_state.username, inst_name, inst_address, inst_contact,
                                  teacher_name, paper_language, board_format, logo_bytes, logo_mimetype):
        st.session_state.inst_defaults = get_institution_defaults(st.session_state.username)
        st.sidebar.success("Saved! These will auto-fill next time you log in.")
    else:
        st.sidebar.error("Couldn't save your defaults. Please try again.")

if inst_logo_upload is not None:
    inst_logo = inst_logo_upload
elif _d.get("default_logo_base64"):
    try:
        inst_logo = StoredLogo(base64.b64decode(_d["default_logo_base64"]), _d.get("default_logo_mimetype") or "image/png")
    except Exception:
        inst_logo = None
else:
    inst_logo = None

# ==========================================
# --- HTTP REQUESTS: GEMINI API CALLS ---
# ==========================================
active_api_key = user_api_key if user_api_key.strip() != "" else SERVER_API_KEY

@st.cache_data(ttl=3600, show_spinner=False)
def get_working_model_name(api_key):
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
    try:
        res = requests.get(url, timeout=30)
        if res.status_code == 200:
            models = [m['name'] for m in res.json().get('models', []) if 'generateContent' in m.get('supportedGenerationMethods', [])]
            flash_models = [m for m in models if '1.5-flash' in m]
            return flash_models[0].replace('models/', '') if flash_models else models[0].replace('models/', '')
        return "gemini-1.5-flash"
    except Exception as e:
        logging.error(f"[Model Fetch Error] {e}")
        return "gemini-1.5-flash"

working_model_name = get_working_model_name(active_api_key)

def generate_gemini_content(prompt, api_key, model_name="gemini-1.5-flash", images=None):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
    headers = {'Content-Type': 'application/json'}
    
    parts = [{"text": prompt}]
    
    if images:
        for img in images:
            buffered = BytesIO()
            # PNGs (and some other formats) can be in RGBA/P/LA mode, which
            # the JPEG encoder can't write (it has no alpha channel support).
            # Converting to RGB first prevents a crash on transparent PNGs.
            if img.mode != "RGB":
                img = img.convert("RGB")
            img.save(buffered, format="JPEG")
            img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
            parts.append({
                "inline_data": {
                    "mime_type": "image/jpeg",
                    "data": img_str
                }
            })
            
    payload = {"contents": [{"parts": parts}]}
    
    # Timeout prevents the app from hanging forever for a user if the
    # Gemini API stalls or the network drops mid-request.
    response = requests.post(url, headers=headers, json=payload, timeout=90)
    if response.status_code == 200:
        data = response.json()
        try:
            return data['candidates'][0]['content']['parts'][0]['text']
        except (KeyError, IndexError):
            logging.error("Unexpected response format from Gemini API.")
            raise Exception("Unexpected response format from Gemini API.")
    else:
        try:
            error_msg = response.json().get('error', {}).get('message', 'Unknown error')
        except ValueError:
            error_msg = response.text[:200]
        logging.error(f"Gemini API Error {response.status_code}: {error_msg}")
        raise Exception(f"API Error {response.status_code}: {error_msg}")

# --- Helper Functions ---
def extract_text_from_pdf(uploaded_file, start_page, end_page):
    try:
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        start_index = max(0, start_page - 1) 
        end_index = min(len(doc), end_page)
        text = ""
        for i in range(start_index, end_index):
            text += doc[i].get_text("text") + "\n"
        return text
    except Exception as e:
        logging.error(f"[PDF Extraction Error] {e}")
        return ""

def render_question_config(key_prefix=""):
    """Renders the Type/Count/Marks/Difficulty grid (MCQ, FIB, True/False,
    Short, Long) and returns all the values. key_prefix keeps widget keys
    unique when this is rendered more than once on the same page (e.g. once
    for the general Create Paper tab, once for the BSEB tab)."""
    h1, h2, h3, h4 = st.columns([3, 2, 2, 3])
    h1.write("**Type**"); h2.write("**Count**"); h3.write("**Marks**"); h4.write("**Diff**")

    c1, c2, c3, c4 = st.columns([3, 2, 2, 3])
    c1.write("MCQs")
    mcq_c = c2.number_input("mcq_c", 0, 50, 5, label_visibility="collapsed", key=f"{key_prefix}m_c")
    mcq_m = c3.number_input("mcq_m", 1, 10, 1, label_visibility="collapsed", key=f"{key_prefix}m_m")
    mcq_d = c4.selectbox("mcq_d", ["Easy", "Medium", "Hard"], label_visibility="collapsed", key=f"{key_prefix}m_d")

    c1, c2, c3, c4 = st.columns([3, 2, 2, 3])
    c1.write("Fill in the Blanks")
    fib_c = c2.number_input("fib_c", 0, 20, 3, label_visibility="collapsed", key=f"{key_prefix}f_c")
    fib_m = c3.number_input("fib_m", 1, 10, 1, label_visibility="collapsed", key=f"{key_prefix}f_m")
    fib_d = c4.selectbox("fib_d", ["Easy", "Medium", "Hard"], label_visibility="collapsed", key=f"{key_prefix}f_d")

    c1, c2, c3, c4 = st.columns([3, 2, 2, 3])
    c1.write("True / False")
    tf_c = c2.number_input("tf_c", 0, 20, 3, label_visibility="collapsed", key=f"{key_prefix}t_c")
    tf_m = c3.number_input("tf_m", 1, 10, 1, label_visibility="collapsed", key=f"{key_prefix}t_m")
    tf_d = c4.selectbox("tf_d", ["Easy", "Medium", "Hard"], label_visibility="collapsed", key=f"{key_prefix}t_d")

    c1, c2, c3, c4 = st.columns([3, 2, 2, 3])
    c1.write("Short Answer")
    short_c = c2.number_input("sh_c", 0, 20, 3, label_visibility="collapsed", key=f"{key_prefix}s_c")
    short_m = c3.number_input("sh_m", 1, 10, 2, label_visibility="collapsed", key=f"{key_prefix}s_m")
    short_d = c4.selectbox("sh_d", ["Easy", "Medium", "Hard"], label_visibility="collapsed", key=f"{key_prefix}s_d", index=1)

    c1, c2, c3, c4 = st.columns([3, 2, 2, 3])
    c1.write("Long Answer")
    long_c = c2.number_input("l_c", 0, 20, 2, label_visibility="collapsed", key=f"{key_prefix}l_c")
    long_m = c3.number_input("l_m", 1, 20, 5, label_visibility="collapsed", key=f"{key_prefix}l_m")
    long_d = c4.selectbox("l_d", ["Easy", "Medium", "Hard"], label_visibility="collapsed", key=f"{key_prefix}l_d", index=2)

    total_q = mcq_c + fib_c + tf_c + short_c + long_c
    total_m = (mcq_c * mcq_m) + (fib_c * fib_m) + (tf_c * tf_m) + (short_c * short_m) + (long_c * long_m)

    st.markdown("---")
    st.info(f"📊 Total Questions: {total_q} | 🏆 Maximum Marks: {total_m}")

    return (mcq_c, mcq_d, mcq_m, fib_c, fib_d, fib_m, tf_c, tf_d, tf_m,
            short_c, short_d, short_m, long_c, long_d, long_m, total_q, total_m)

def build_question_prompt(mcq_c, mcq_d, mcq_m, fib_c, fib_d, fib_m, tf_c, tf_d, tf_m, short_c, short_d, short_m, long_c, long_d, long_m, include_answers, selected_language, subject):
    reqs = []
    if mcq_c > 0: reqs.append(f"## Multiple Choice Questions [{mcq_m} Mark(s) Each]\n- {mcq_c} MCQs (Difficulty: {mcq_d}).")
    if fib_c > 0: reqs.append(f"## Fill in the Blanks [{fib_m} Mark(s) Each]\n- {fib_c} FIBs (Difficulty: {fib_d}). MUST include 4 options (A, B, C, D) on a new line for each blank.")
    if tf_c > 0:  reqs.append(f"## True / False [{tf_m} Mark(s) Each]\n- {tf_c} True/False questions (Difficulty: {tf_d}). MUST include exactly 2 options: (A) True  (B) False on a new line.")
    if short_c > 0: reqs.append(f"## Short Answer Questions [{short_m} Mark(s) Each]\n- {short_c} Short Qs (Difficulty: {short_d}).")
    if long_c > 0:  reqs.append(f"## Long Answer Questions [{long_m} Mark(s) Each]\n- {long_c} Long Qs (Difficulty: {long_d}).")
    
    if selected_language == "English":
        lang_instruction = "LANGUAGE RULE: Generate the ENTIRE paper and answers strictly in the English language."
    elif selected_language == "Hindi":
        lang_instruction = "LANGUAGE RULE: Generate the paper in simple Hindi. Avoid tough academic Hindi words. Provide English terms in brackets for technical words. Example: 'अंश [Numerator]'."
    else:
        lang_instruction = "LANGUAGE RULE: Generate the paper in Hinglish (a mix of simple Hindi and English). Provide English terms in brackets for technical words."
    
    base_prompt = "\n\n".join(reqs) + f"\n\n{lang_instruction}\n\n" + f"""CRITICAL FORMATTING:
1. STRICTLY adhere to the subject: **{subject}**. Do NOT generate general knowledge questions.
2. CONTINUOUS NUMBERING: Number ALL questions continuously from start to finish (e.g., **Q1.**, **Q2.**, **Q3.**, etc.) across ALL sections. Do NOT restart numbering at 1 for a new section. 
3. OPTIONS ON NEW LINE: For MCQs, FIBs, and T/F, ALWAYS place the options on a NEW LINE directly below the question text. Do NOT place options on the same line as the question.
   Correct Example:
   **Q1.** What is the value of x?
   (A) 1   (B) 2   (C) 3   (D) 4
4. MARKS IN HEADERS: Include the marks per question in the section headers as provided above.
5. DETAILED ANSWERS: In the Answer Key, provide detailed, step-by-step explanations for Short and Long answer questions, proportional to their marks (e.g., 5-mark questions need a long, detailed explanation). Ensure answer numbers match the continuous question numbers.
6. DELIMITER: Separate EVERY single Question, Section Header, and the Answer Key with the delimiter `|||` on a new line. Do not group multiple questions together.
7. MATH: USE UNICODE SYMBOLS ONLY (θ, π, √, ²). NO LaTeX. Write fractions as a/b.
8. DO NOT generate any Title, Institute Name, Time, or Marks at the very top. Start directly with the first section header.
    """
    
    if include_answers: return base_prompt + "\nAdd '# ANSWER KEY' at end, also separated by `|||`. Ensure numbering in answers exactly matches the continuous numbering of the questions."
    return base_prompt

def extract_question_number(text):
    """Pulls the question number out of a block like '**Q3.** ...' or '3. ...'
    so we can find the matching Answer Key entry for a regenerated question."""
    m = re.search(r'Q\.?\s*(\d+)', text, re.IGNORECASE)
    if m:
        return m.group(1)
    m2 = re.match(r'\s*\**\s*(\d+)\.', text)
    if m2:
        return m2.group(1)
    return None

def regenerate_single_question(old_text, api_key, model_name, subject, topics):
    """Regenerates one question, staying on-topic, and also returns a fresh
    answer/solution for it so the Answer Key can be kept in sync."""
    topic_context = topics.strip() if topics and topics.strip() else subject
    prompt = (
        f"You are regenerating ONE question from a {subject} exam paper. "
        f"The paper's topics are: {topic_context}. Stay strictly within this subject and these topics — "
        f"do not drift into unrelated topics.\n\n"
        f"The original question being replaced was:\n{old_text}\n\n"
        "Write a NEW question of the same type, difficulty, and marks as the original, strictly on the same "
        "subject/topics. Keep the same question number label if the original had one (e.g. 'Q3.'). "
        "Use Unicode math symbols (θ, π, √, ²) instead of LaTeX.\n\n"
        "Then on a new line write the exact delimiter @@@ANSWER@@@ followed by the correct answer/solution for "
        "THIS NEW question — a brief correct option letter for MCQ/True-False/Fill-in-the-blank, or a full "
        "step-by-step explanation for Short/Long answer questions.\n\n"
        "Output ONLY the question text, then @@@ANSWER@@@, then the answer. No extra commentary."
    )
    resp_text = generate_gemini_content(prompt, api_key, model_name)
    if "@@@ANSWER@@@" in resp_text:
        q_part, a_part = resp_text.split("@@@ANSWER@@@", 1)
        return q_part.strip(), a_part.strip()
    return resp_text.strip(), None

def clean_math_for_word(text):
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\\\((.*?)\\\)', r'\1', text)
    text = re.sub(r'\\\[(.*?)\\\]', r'\1', text)
    latex_map = {r'\pi': 'π', r'\theta': 'θ', r'\sqrt': '√', r'\times': '×', r'\div': '÷', '$': '', '^2': '²', '^3': '³'}
    for k, v in latex_map.items(): text = text.replace(k, v)
    text = text.replace('☐', '[ ]').replace('☑', '[x]').replace('•', '-').replace('◦', '-')
    text = text.replace('\u200b', '').replace('\u2022', '-').replace('\u25cf', '-').replace('\u25cb', '-')
    return text.strip()

# 🌟 HTML RENDERER 🌟
def create_a4_html(md_content, i_name, i_address, i_contact, t_name, inst_logo=None, is_2_col=False, sub="Subject", grade="Class", total_m="Marks", exam_time="Time", topics=""):
    md_content = clean_math_for_word(md_content)
    
    md_content = re.sub(r"^#.*?\*\*\*", "", md_content, count=1, flags=re.DOTALL).strip()
    md_content = re.sub(r"^\*\*Subject:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Class:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Marks:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Time:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    
    md_content = re.sub(r"^\d+\.\s", "**Q.** ", md_content, flags=re.MULTILINE)
    md_content = md_content.strip()
    
    logo_html_inline = ""
    logo_footer = ""
    if inst_logo:
        inst_logo.seek(0)
        b64 = base64.b64encode(inst_logo.getvalue()).decode()
        logo_html_inline = f"<td style='width: 1%; padding-right: 15px; vertical-align: middle;'><img src='data:{inst_logo.type};base64,{b64}' style='max-height: 55px;'/></td>"
        logo_footer = f"<img src='data:{inst_logo.type};base64,{b64}' style='height: 18px; vertical-align: middle; margin-right: 8px;'/>"
    
    main_heading_text = topics.strip().upper() if topics.strip() != "" else sub.upper()
    
    custom_header = f"""
    <div style='border-bottom: 2px solid black; padding-bottom: 10px; margin-bottom: 10px; width: 100%;'>
        <table style='width: 100%; border-collapse: collapse; border: none; margin-bottom: 10px;'>
            <tr>
                <td style='text-align: center; vertical-align: middle; border: none;'>
                    <table style='margin: 0 auto;'>
                        <tr>
                            {logo_html_inline}
                            <td style='vertical-align: middle;'>
                                <h1 style='margin: 0; font-size: 24px; font-family: "Noto Sans", "Nirmala UI", "Times New Roman", serif; font-weight: 900; text-transform: uppercase; white-space: nowrap;'>{i_name}</h1>
                            </td>
                        </tr>
                    </table>
                </td>
            </tr>
        </table>
        <table style='width: 100%; font-weight: bold; font-size: 13px; border: none;'>
            <tr>
                <td style='text-align: left; vertical-align: bottom; width: 33%; border: none;'>Class : {grade}<br>Time : {exam_time}</td>
                <td style='text-align: center; vertical-align: middle; width: 34%; border: none;'>
                    <div style='border: 2px solid black; border-radius: 12px; display: inline-block; padding: 4px 25px; font-weight: bold; font-size: 14px; background: white;'>
                        EXAMINATION
                    </div>
                </td>
                <td style='text-align: right; vertical-align: bottom; width: 33%; border: none;'>Sub.: {sub}<br>Marks: {total_m}</td>
            </tr>
        </table>
    </div>
    <div style='border-top: 1px solid black; border-bottom: 3px solid black; padding: 2px 0; margin-bottom: 15px;'>
        <div style='background-color: black; color: white; padding: 5px; text-align: center; font-weight: bold; font-size: 15px; text-transform: uppercase; letter-spacing: 1px;'>
            Multiple Choice Questions & Theory
        </div>
    </div>
    <h2 style='text-align: center; text-decoration: underline; text-transform: uppercase; margin-top: 0; margin-bottom: 15px; font-size: 18px;'>{main_heading_text}</h2>
    """

    ans_split_marker = "|||ANSWER_KEY_SPLIT|||"
    md_content = re.sub(r'(?im)^#+\s*Answer Key.*$', ans_split_marker, md_content)
    
    if ans_split_marker in md_content:
        q_part, a_part = md_content.split(ans_split_marker)
        final_inner_html = f"""
        {custom_header}
        <div class="content-body">{markdown.markdown(q_part.strip())}</div>
        <div style="page-break-before: always; width: 100%;"></div>
        {custom_header}
        <h2 style="text-align: center; text-decoration: underline; margin-bottom: 15px;">ANSWER KEY</h2>
        <div class="content-body">{markdown.markdown(a_part.strip())}</div>
        """
    else:
        final_inner_html = f"""
        {custom_header}
        <div class="content-body">{markdown.markdown(md_content.strip())}</div>
        """
    
    col_style = "column-count: 2; column-gap: 15mm; column-rule: 1px solid #000; font-size: 14px;" if is_2_col else "font-size: 16px;"

    return f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><style>
    body {{ background: #f0f0f0; font-family: 'Noto Sans', 'Nirmala UI', 'Times New Roman', serif; margin: 0; padding: 20px; display: flex; justify-content: center; }} 
    .a4-page {{ background: white; width: 210mm; min-height: 297mm; padding: 20px; box-shadow: 0 0 10px rgba(0,0,0,0.2); box-sizing: border-box; position: relative; overflow: hidden; }} 
    .watermark {{ position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%) rotate(-45deg); font-size: 85px; color: rgba(0, 0, 0, 0.06); z-index: -9999; pointer-events: none; white-space: nowrap; font-weight: bold; text-transform: uppercase; }}
    table {{ width: 100%; border-collapse: collapse; border: none; position: relative; z-index: 1; }}
    td {{ border: none; padding: 0; }}
    @media print {{ 
        @page {{ size: A4; margin: 0; }} 
        body {{ background: white; padding: 0; margin: 0; display: block; }} 
        .a4-page {{ box-shadow: none; width: 100%; min-height: auto; padding: 10mm; margin: 0; page-break-after: always; }} 
        .watermark {{ color: rgba(0, 0, 0, 0.06) !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
        tfoot {{ display: table-footer-group; }}
    }} 
    h1, h2, h3 {{ text-align: center; column-span: all; }} 
    h2 {{ font-size: 16px; border-bottom: 1px dashed #ccc; padding-bottom: 5px; }}
    .content-body {{ {col_style} position: relative; z-index: 1; text-align: justify; }} 
    .content-body p {{ margin-bottom: 8px; margin-top: 4px; }}
    .footer-content {{ text-align: center; margin-top: 20px; padding-top: 10px; border-top: 2px dashed #bbb; font-size: 13px; color: #444; position: relative; z-index: 1; background: white; }}
    </style></head><body><div class="a4-page">
    <div class="watermark">{i_name}</div>
    <table>
        <thead><tr><td></td></tr></thead>
        <tbody><tr><td>{final_inner_html}</td></tr></tbody>
        <tfoot><tr><td>
            <div class="footer-content">
                {logo_footer}<strong>{i_name}</strong> | 📍 {i_address} | 📞 {i_contact} | 👨‍🏫 <strong>{t_name}</strong>
            </div>
        </td></tr></tfoot>
    </table>
    </div></body></html>"""

def html_to_pdf(html_string):
    try:
        buf = BytesIO()
        result = pisa.CreatePDF(html_string, dest=buf)
        if result.err:
            return None
        return buf.getvalue()
    except Exception as e:
        logging.error(f"[PDF Generation Error] {e}")
        return None

# 🌟 WORD RENDERER 🌟
def create_word_docx(md_content, i_name, i_address, i_contact, t_name, inst_logo=None, is_2_col=False, sub="Subject", grade="Class", total_m="Marks", exam_time="Time", topics=""):
    doc = Document()
    
    md_content = re.sub(r"^#.*?\*\*\*", "", md_content, count=1, flags=re.DOTALL).strip()
    md_content = re.sub(r"^\*\*Subject:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Class:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Marks:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\*\*Time:\*\*.*?\n", "", md_content, flags=re.MULTILINE)
    md_content = re.sub(r"^\d+\.\s", "**Q.** ", md_content, flags=re.MULTILINE)
    md_content = md_content.strip()
        
    md_content = md_content.replace('\r', '')
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial' 
    font.size = Pt(11)
    
    rFonts = style.element.rPr.rFonts
    if rFonts is not None:
        rFonts.set(qn('w:cs'), 'Noto Sans Devanagari') 
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
    style_lang = style.element.rPr.find(qn('w:lang'))
    if style_lang is None:
        style_lang = style.element.rPr.makeelement(qn('w:lang'), {})
        style.element.rPr.append(style_lang)
    style_lang.set(qn('w:bidi'), 'hi-IN')
    
    for i in range(3):
        try:
            h_style = doc.styles[f'Heading {i}']
            h_style.font.name = 'Arial'
            if h_style.element.rPr.rFonts is not None:
                h_style.element.rPr.rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
                h_style.element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                h_style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            if i == 0:
                h_style.font.size = Pt(16)
                h_style.font.bold = True
            elif i == 1:
                h_style.font.size = Pt(12)
                h_style.font.bold = True
            elif i == 2:
                h_style.font.size = Pt(11)
                h_style.font.bold = True
        except KeyError: pass

    if is_2_col:
        for section in doc.sections:
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.4)

    def apply_cs_font(run):
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.append(rfonts)
        rfonts.set(qn('w:cs'), 'Noto Sans Devanagari')
        rfonts.set(qn('w:ascii'), 'Arial')
        rfonts.set(qn('w:hAnsi'), 'Arial')
        # Without an explicit language tag, Word doesn't reliably classify
        # Devanagari text as "complex script" and may render it with the
        # ascii font (Arial, no Devanagari glyphs = tofu boxes) regardless
        # of the w:cs font specified above. This tag is what makes Word
        # actually route the text correctly.
        lang = rpr.find(qn('w:lang'))
        if lang is None:
            lang = rpr.makeelement(qn('w:lang'), {})
            rpr.append(lang)
        lang.set(qn('w:bidi'), 'hi-IN')

    def insert_chate_header():
        title_table = doc.add_table(rows=1, cols=1)
        p1 = title_table.cell(0,0).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if inst_logo is not None:
            try:
                inst_logo.seek(0)
                r_logo = p1.add_run()
                r_logo.add_picture(inst_logo, height=Inches(0.38))
                p1.add_run("   ") 
            except Exception: pass
            
        r1 = p1.add_run(i_name.upper())
        r1.bold = True
        r1.font.size = Pt(18)
        apply_cs_font(r1)
        
        details_table = doc.add_table(rows=1, cols=3)
        details_table.autofit = False
        for cell in details_table.columns[0].cells: cell.width = Inches(2.0)
        for cell in details_table.columns[1].cells: cell.width = Inches(3.0)
        for cell in details_table.columns[2].cells: cell.width = Inches(2.0)

        p3 = details_table.cell(0,0).paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r3 = p3.add_run(f"Class : {grade}\nTime : {exam_time}")
        r3.bold = True
        r3.font.size = Pt(10)
        apply_cs_font(r3)

        p4 = details_table.cell(0,1).paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run("\n[ EXAMINATION ]")
        r4.bold = True
        r4.font.size = Pt(12)
        apply_cs_font(r4)

        p2 = details_table.cell(0,2).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(f"Sub.: {sub}\nMarks: {total_m}")
        r2.bold = True
        r2.font.size = Pt(10)
        apply_cs_font(r2)
        
        doc.add_paragraph("__________________________________________________________________________").alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt = doc.add_paragraph("MULTIPLE CHOICE QUESTIONS & THEORY")
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.runs[0].bold = True
        apply_cs_font(pt.runs[0])
        
        main_heading_text = topics.strip().upper() if topics.strip() != "" else sub.upper()
        ptopics = doc.add_paragraph(main_heading_text)
        ptopics.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ptopics.runs[0].underline = True
        ptopics.runs[0].font.size = Pt(14)
        ptopics.runs[0].bold = True
        apply_cs_font(ptopics.runs[0])
        doc.add_paragraph() 

    insert_chate_header()

    if is_2_col:
        new_section = doc.add_section(0) 
        sectPr = new_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720') 

    for line in md_content.split('\n'):
        line_clean = line.strip()
        if not line_clean: continue
        line_clean = clean_math_for_word(line_clean)
        
        if "Answer Key" in line_clean or "ANSWER KEY" in line_clean:
            doc.add_page_break() 
            insert_chate_header() 
            doc.add_heading("Answer Key", level=1)
            continue
            
        if line_clean.startswith('# '): 
            doc.add_heading(line_clean.replace('# ', ''), level=1)
        elif line_clean.startswith('## '): 
            doc.add_heading(line_clean.replace('## ', ''), level=2)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            parts = re.split(r'\*\*(.*?)\*\*', line_clean)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: run.bold = True
                apply_cs_font(run)
                
    if doc.sections:
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if inst_logo is not None:
            try:
                inst_logo.seek(0)
                run_logo = footer_para.add_run()
                run_logo.add_picture(inst_logo, height=Inches(0.18))
                footer_para.add_run("  ") 
            except Exception: pass
            
        run_name = footer_para.add_run(f"{i_name}  |  ")
        run_name.font.size = Pt(10)
        run_name.font.bold = True
        run_name.font.color.rgb = RGBColor(100, 100, 100)
        apply_cs_font(run_name)
        
        run_rest = footer_para.add_run(f"📍 {i_address}  |  📞 {i_contact}  |  👨‍🏫 {t_name}")
        run_rest.font.size = Pt(10)
        run_rest.font.color.rgb = RGBColor(100, 100, 100)
        apply_cs_font(run_rest)
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# --- MAIN LAYOUT ---
# ==========================================
tab_create, tab_bseb, tab_digitize, tab_history = st.tabs(["🏠 Create Paper", "🎓 BSEB Board", "📷 Digitize Handwritten", "🗂️ Cloud History"])

with tab_create:
    if not is_pro and papers_used >= FREE_LIMIT:
        st.error("Free Trial Expired! Please Upgrade.")
        st.stop()
        
    st.markdown("### 1. Source")
    source = st.radio("Method:", ["⚡ Quick", "📄 PDF Extract"], horizontal=True, label_visibility="collapsed")

    sub, grade, syl, up_pdf = "", "", "", None
    pdf_text = ""

    if source == "⚡ Quick":
        c1, c2 = st.columns(2)
        sub = c1.text_input("Subject")
        grade = c2.text_input("Class")
        syl = st.text_area("Topics")
    else:
        c1, c2 = st.columns(2)
        sub = c1.text_input("Subject (PDF)")
        grade = c2.text_input("Class (PDF)")
        syl = st.text_area("Specific Topics (Optional)")
        
        up_pdf = st.file_uploader("Upload PDF Book/Notes", type="pdf")
        
        c3, c4 = st.columns(2)
        start_p = c3.number_input("Start Page", min_value=1, value=1)
        end_p = c4.number_input("End Page", min_value=1, value=5)
        
        if up_pdf is not None:
            pdf_text = extract_text_from_pdf(up_pdf, start_p, end_p)
            st.success(f"Extracted {len(pdf_text)} characters from pages {start_p} to {end_p}.")

    st.markdown("---")
    st.markdown("### 2. Counts & Marks")
    (mcq_c, mcq_d, mcq_m, fib_c, fib_d, fib_m, tf_c, tf_d, tf_m,
     short_c, short_d, short_m, long_c, long_d, long_m, total_q, total_m) = render_question_config()

    if "blocks_saved" not in st.session_state: st.session_state.blocks_saved = True
    if "confirm_overwrite" not in st.session_state: st.session_state.confirm_overwrite = False

    generate_clicked = st.button("🚀 Generate Paper", use_container_width=True)

    if generate_clicked and st.session_state.blocks and not st.session_state.blocks_saved and not st.session_state.confirm_overwrite:
        st.warning("You have an unsaved paper above (not saved to Cloud History yet). Generating a new one will replace it.")
        if st.button("Generate anyway (discard current paper)", use_container_width=True):
            st.session_state.confirm_overwrite = True
            st.rerun()
        generate_clicked = False

    if generate_clicked and (not st.session_state.blocks or st.session_state.blocks_saved or st.session_state.confirm_overwrite):
        st.session_state.confirm_overwrite = False
        if not sub.strip() or not grade.strip():
            st.error("Please fill in Subject and Class before generating.")
        elif total_q == 0:
            st.error("Please add at least one question (set a count > 0 for some question type).")
        else:
            st.session_state.current_subject = sub
            st.session_state.current_class = grade
            st.session_state.current_marks = str(total_m)

            q_reqs = build_question_prompt(
                mcq_c, mcq_d, mcq_m, fib_c, fib_d, fib_m, tf_c, tf_d, tf_m, short_c, short_d, short_m, long_c, long_d, long_m, include_answer_key, paper_language, sub
            )

            if source == "📄 PDF Extract" and pdf_text != "":
                prompt = f"Subject: {sub}\nClass: {grade}\nTopics: {syl}\n\n{q_reqs}\n\nIMPORTANT: Start directly with the questions. DO NOT generate any Title, Institute Name, Time, or Marks at the top.\n\nCREATE QUESTIONS STRICTLY FROM THE FOLLOWING TEXT EXTRACTED FROM A BOOK:\n\n{pdf_text}"
            else:
                prompt = f"Subject: {sub}\nClass: {grade}\nTopics: {syl}\n\n{q_reqs}\n\nIMPORTANT: Start directly with the questions. DO NOT generate any Title, Institute Name, Time, or Marks at the top."

            with st.spinner("Generating Paper..."):
                try:
                    resp_text = generate_gemini_content(prompt, active_api_key, working_model_name)
                    blocks = resp_text.split("|||")
                    st.session_state.blocks = [{'id': str(uuid.uuid4()), 'text': b.strip()} for b in blocks if b.strip()]
                    st.session_state.blocks_saved = False
                    st.session_state.file_name = f"{sub}_Paper"
                    update_paper_count(st.session_state.username)
                    st.rerun()
                except Exception as e:
                    error_msg = str(e).lower()
                    logging.error(f"[Generation Error] {e}") 
                    if "429" in error_msg or "quota" in error_msg:
                        st.error("🚨 The daily generation limit has been reached! Try again later or add your own API key in Advanced Settings.")
                    else:
                        st.error("Something went wrong generating the paper. Please try again.")

    if st.session_state.blocks:
        st.markdown("---")
        with st.expander("🛠️ Edit Questions", expanded=False):
            st.caption("Regenerating a question also updates its matching Answer Key entry, if one exists.")
            for i, b in enumerate(st.session_state.blocks):
                edit_col, regen_col = st.columns([5, 1])
                new_text = edit_col.text_area(f"Question {i+1}", b['text'], height=100, key=f"block_text_{b['id']}")
                if new_text != st.session_state.blocks[i]['text']:
                    st.session_state.blocks[i]['text'] = new_text
                    st.session_state.blocks_saved = False
                if regen_col.button("🔄 Regenerate", key=f"regen_{b['id']}", help="Ask AI to write a fresh version of this question (and its answer key entry, if present)"):
                    with st.spinner("Regenerating..."):
                        try:
                            old_number = extract_question_number(b['text'])
                            new_q_text, new_answer_text = regenerate_single_question(b['text'], active_api_key, working_model_name, sub, syl)
                            st.session_state.blocks[i]['text'] = new_q_text
                            st.session_state.blocks[i]['id'] = str(uuid.uuid4())

                            # Keep the Answer Key in sync: find the matching answer
                            # entry (by question number) and update it too, so the
                            # solution doesn't stay pointing at the old question.
                            if new_answer_text and old_number:
                                ans_key_idx = None
                                for j, blk in enumerate(st.session_state.blocks):
                                    if "ANSWER KEY" in blk['text'].upper():
                                        ans_key_idx = j
                                        break
                                if ans_key_idx is not None:
                                    for j in range(ans_key_idx + 1, len(st.session_state.blocks)):
                                        blk_number = extract_question_number(st.session_state.blocks[j]['text'])
                                        if blk_number == old_number:
                                            st.session_state.blocks[j]['text'] = f"**Q{old_number}.** {new_answer_text}"
                                            st.session_state.blocks[j]['id'] = str(uuid.uuid4())
                                            break

                            st.session_state.blocks_saved = False
                            st.rerun()
                        except Exception as e:
                            logging.error(f"[Regenerate Error] {e}")
                            st.error("Couldn't regenerate this question. Please try again.")
        
        paper_md = "\n\n".join([b['text'] for b in st.session_state.blocks])
        
        f_html = create_a4_html(paper_md, inst_name, inst_address, inst_contact, teacher_name, inst_logo, is_two_column, st.session_state.current_subject, st.session_state.current_class, st.session_state.current_marks, exam_time, syl)
        f_word = create_word_docx(paper_md, inst_name, inst_address, inst_contact, teacher_name, inst_logo, is_two_column, st.session_state.current_subject, st.session_state.current_class, st.session_state.current_marks, exam_time, syl)
        
        c1, c2, c3, c4 = st.columns(4)
        c1.download_button("🖨️ HTML", f_html, f"{st.session_state.current_subject}.html", "text/html")
        c2.download_button("📄 Word", f_word, f"{st.session_state.current_subject}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        f_pdf = html_to_pdf(f_html)
        if f_pdf:
            c3.download_button("📕 PDF", f_pdf, f"{st.session_state.current_subject}.pdf", "application/pdf")
        else:
            c3.caption("PDF unavailable")
        if paper_language in ("Hindi", "Bilingual"):
            st.caption("💡 Hindi text in the Word file needs the free 'Noto Sans Devanagari' font installed on the computer opening it (one-time setup). It's not needed for the PDF.")
        if c4.button("☁️ Save History"):
            data = {"username": st.session_state.username, "date": datetime.now().strftime("%Y-%m-%d"), "subject": st.session_state.current_subject, "board": board_format, "content": paper_md}
            try:
                supabase.table("papers").insert(data).execute()
                st.session_state.blocks_saved = True
                st.success("Saved!")
            except Exception as e:
                logging.error(f"[Save History Error] {e}")
                st.error("Couldn't save to Cloud History. Please try again.")

with tab_bseb:
    if not is_pro and papers_used >= FREE_LIMIT:
        st.error("Free Trial Expired! Please Upgrade.")
        st.stop()

    st.markdown("### 🎓 Bihar Board (BSEB) Paper Builder")
    st.caption("Pick class, subject(s), and chapters from your saved syllabus list — then set question types and marks. Tip: set Board Pattern to 'BSEB (Bihar Board)' in the sidebar for BSEB-style formatting.")

    st.markdown("#### 1. Class & Subjects")
    bc1, bc2 = st.columns(2)
    bseb_class = bc1.selectbox("Class", CLASS_OPTIONS, key="bseb_class")
    bseb_existing_subjects = get_subjects_for_class(bseb_class)
    bseb_sel_subjects = bc2.multiselect("Subject(s)", bseb_existing_subjects, key="bseb_subjects")

    bseb_new_subject = st.text_input("Add a new subject for this class (optional)", key="bseb_new_subject", placeholder="e.g. Mathematics")
    if bseb_new_subject.strip() and bseb_new_subject.strip() not in bseb_sel_subjects:
        bseb_sel_subjects = bseb_sel_subjects + [bseb_new_subject.strip()]

    st.markdown("#### 2. Chapters")
    bseb_all_chapters = []
    for subj in bseb_sel_subjects:
        with st.expander(f"📖 {subj} — chapters", expanded=True):
            existing_chapters = get_chapters(bseb_class, subj)
            chosen_chapters = st.multiselect(f"Select chapters ({subj})", existing_chapters, key=f"bseb_chap_sel_{bseb_class}_{subj}")
            add_chapters_text = st.text_area(
                f"Add new chapters for {subj} (comma-separated)", key=f"bseb_chap_add_{bseb_class}_{subj}",
                placeholder="e.g. Real Numbers, Polynomials, Triangles",
                help="Saved for everyone using PaperBanao — next time, just select them instead of retyping."
            )
            if st.button(f"💾 Save chapters for {subj}", key=f"bseb_chap_save_{bseb_class}_{subj}"):
                new_list = existing_chapters + [c.strip() for c in add_chapters_text.split(",") if c.strip()]
                if save_chapters(bseb_class, subj, new_list):
                    st.success(f"Saved! Chapters for {bseb_class} - {subj} updated.")
                    st.rerun()
                else:
                    st.error("Couldn't save chapters. Please try again.")
            bseb_all_chapters.extend([f"{subj}: {c}" for c in chosen_chapters])
            bseb_all_chapters.extend([f"{subj}: {c.strip()}" for c in add_chapters_text.split(",") if c.strip()])

    bseb_chapter_topics = "; ".join(bseb_all_chapters)
    if bseb_chapter_topics:
        st.caption(f"✅ Selected: {bseb_chapter_topics}")

    st.markdown("#### 3. Specific Topics (optional)")
    bseb_extra_topics = st.text_area("Narrow it down further within the chosen chapters, if needed", key="bseb_extra_topics")
    bseb_syl = ", ".join(filter(None, [bseb_chapter_topics, bseb_extra_topics.strip()]))
    bseb_sub = ", ".join(bseb_sel_subjects) if bseb_sel_subjects else ""

    st.markdown("---")
    st.markdown("#### 4. Counts & Marks")
    (b_mcq_c, b_mcq_d, b_mcq_m, b_fib_c, b_fib_d, b_fib_m, b_tf_c, b_tf_d, b_tf_m,
     b_short_c, b_short_d, b_short_m, b_long_c, b_long_d, b_long_m, b_total_q, b_total_m) = render_question_config(key_prefix="bseb_")

    if "bseb_blocks" not in st.session_state: st.session_state.bseb_blocks = []
    if "bseb_blocks_saved" not in st.session_state: st.session_state.bseb_blocks_saved = True

    if st.button("🚀 Generate BSEB Paper", use_container_width=True):
        if not bseb_sel_subjects:
            st.error("Please select at least one subject.")
        elif not bseb_chapter_topics and not bseb_extra_topics.strip():
            st.error("Please select at least one chapter, or add specific topics.")
        elif b_total_q == 0:
            st.error("Please add at least one question (set a count > 0 for some question type).")
        else:
            st.session_state.current_subject = bseb_sub
            st.session_state.current_class = bseb_class
            st.session_state.current_marks = str(b_total_m)

            b_q_reqs = build_question_prompt(
                b_mcq_c, b_mcq_d, b_mcq_m, b_fib_c, b_fib_d, b_fib_m, b_tf_c, b_tf_d, b_tf_m,
                b_short_c, b_short_d, b_short_m, b_long_c, b_long_d, b_long_m,
                include_answer_key, paper_language, bseb_sub
            )
            b_prompt = (
                f"Subject: {bseb_sub}\nClass: {bseb_class}\nBoard: Bihar Board (BSEB)\nTopics: {bseb_syl}\n\n{b_q_reqs}\n\n"
                "IMPORTANT: Start directly with the questions. DO NOT generate any Title, Institute Name, Time, or Marks at the top."
            )

            with st.spinner("Generating BSEB Paper..."):
                try:
                    resp_text = generate_gemini_content(b_prompt, active_api_key, working_model_name)
                    b_blocks = resp_text.split("|||")
                    st.session_state.bseb_blocks = [{'id': str(uuid.uuid4()), 'text': b.strip()} for b in b_blocks if b.strip()]
                    st.session_state.bseb_blocks_saved = False
                    update_paper_count(st.session_state.username)
                    st.rerun()
                except Exception as e:
                    error_msg = str(e).lower()
                    logging.error(f"[BSEB Generation Error] {e}")
                    if "429" in error_msg or "quota" in error_msg:
                        st.error("🚨 The daily generation limit has been reached! Try again later or add your own API key in Advanced Settings.")
                    else:
                        st.error("Something went wrong generating the paper. Please try again.")

    if st.session_state.bseb_blocks:
        st.markdown("---")
        with st.expander("🛠️ Edit Questions", expanded=False):
            st.caption("Regenerating a question also updates its matching Answer Key entry, if one exists.")
            for i, b in enumerate(st.session_state.bseb_blocks):
                edit_col, regen_col = st.columns([5, 1])
                new_text = edit_col.text_area(f"Question {i+1}", b['text'], height=100, key=f"bseb_block_text_{b['id']}")
                if new_text != st.session_state.bseb_blocks[i]['text']:
                    st.session_state.bseb_blocks[i]['text'] = new_text
                    st.session_state.bseb_blocks_saved = False
                if regen_col.button("🔄 Regenerate", key=f"bseb_regen_{b['id']}", help="Ask AI to write a fresh version of this question (and its answer key entry, if present)"):
                    with st.spinner("Regenerating..."):
                        try:
                            old_number = extract_question_number(b['text'])
                            new_q_text, new_answer_text = regenerate_single_question(b['text'], active_api_key, working_model_name, bseb_sub, bseb_syl)
                            st.session_state.bseb_blocks[i]['text'] = new_q_text
                            st.session_state.bseb_blocks[i]['id'] = str(uuid.uuid4())
                            if new_answer_text and old_number:
                                ans_key_idx = None
                                for j, blk in enumerate(st.session_state.bseb_blocks):
                                    if "ANSWER KEY" in blk['text'].upper():
                                        ans_key_idx = j
                                        break
                                if ans_key_idx is not None:
                                    for j in range(ans_key_idx + 1, len(st.session_state.bseb_blocks)):
                                        blk_number = extract_question_number(st.session_state.bseb_blocks[j]['text'])
                                        if blk_number == old_number:
                                            st.session_state.bseb_blocks[j]['text'] = f"**Q{old_number}.** {new_answer_text}"
                                            st.session_state.bseb_blocks[j]['id'] = str(uuid.uuid4())
                                            break
                            st.session_state.bseb_blocks_saved = False
                            st.rerun()
                        except Exception as e:
                            logging.error(f"[BSEB Regenerate Error] {e}")
                            st.error("Couldn't regenerate this question. Please try again.")

        bseb_paper_md = "\n\n".join([b['text'] for b in st.session_state.bseb_blocks])
        bseb_html = create_a4_html(bseb_paper_md, inst_name, inst_address, inst_contact, teacher_name, inst_logo, is_two_column, bseb_sub, bseb_class, str(b_total_m), exam_time, bseb_syl)
        bseb_word = create_word_docx(bseb_paper_md, inst_name, inst_address, inst_contact, teacher_name, inst_logo, is_two_column, bseb_sub, bseb_class, str(b_total_m), exam_time, bseb_syl)

        bc1, bc2, bc3, bc4 = st.columns(4)
        bc1.download_button("🖨️ HTML", bseb_html, f"{bseb_sub or 'BSEB'}_Paper.html", "text/html")
        bc2.download_button("📄 Word", bseb_word, f"{bseb_sub or 'BSEB'}_Paper.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        bseb_pdf = html_to_pdf(bseb_html)
        if bseb_pdf:
            bc3.download_button("📕 PDF", bseb_pdf, f"{bseb_sub or 'BSEB'}_Paper.pdf", "application/pdf")
        else:
            bc3.caption("PDF unavailable")
        if paper_language in ("Hindi", "Bilingual"):
            st.caption("💡 Hindi text in the Word file needs the free 'Noto Sans Devanagari' font installed on the computer opening it (one-time setup). It's not needed for the PDF.")
        if bc4.button("☁️ Save History", key="bseb_save_history"):
            data = {"username": st.session_state.username, "date": datetime.now().strftime("%Y-%m-%d"), "subject": bseb_sub or "BSEB Paper", "board": "BSEB", "content": bseb_paper_md}
            try:
                supabase.table("papers").insert(data).execute()
                st.session_state.bseb_blocks_saved = True
                st.success("Saved!")
            except Exception as e:
                logging.error(f"[BSEB Save History Error] {e}")
                st.error("Couldn't save to Cloud History. Please try again.")

with tab_digitize:
    st.markdown("### 📷 Digitize a Handwritten Paper")
    st.caption("Upload photos of a handwritten or scanned question paper — we'll read it and turn it into a clean, formatted digital paper using your saved institute details.")

    if "digi_blocks" not in st.session_state: st.session_state.digi_blocks = []
    if "digi_saved" not in st.session_state: st.session_state.digi_saved = True

    digi_images = st.file_uploader("Upload page photos (one or more, in order)", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="digi_uploader")
    dcol1, dcol2 = st.columns(2)
    digi_subject = dcol1.text_input("Subject", key="digi_subject", placeholder="e.g. Mathematics")
    digi_class = dcol2.text_input("Class", key="digi_class", placeholder="e.g. Class 10")

    if st.button("📷 Digitize Paper", use_container_width=True):
        if not digi_images:
            st.error("Please upload at least one photo of the paper.")
        elif not is_pro and papers_used >= FREE_LIMIT:
            st.error("⚠️ Free Trial Expired! Upgrade to Pro from the sidebar to keep generating/digitizing papers.")
        else:
            with st.spinner("Reading your paper... this can take a moment for multiple pages."):
                try:
                    digi_prompt = (
                        "You are digitizing a handwritten or scanned question paper from the attached image(s). "
                        "Transcribe it faithfully — preserve the original question numbering, order, sections, and "
                        "marks exactly as written. Do not invent new questions, do not change the meaning, and do not "
                        "add a title, institute name, header, or footer. Only fix obvious spelling/OCR mistakes. "
                        "Separate each distinct question with the delimiter ||| on its own line."
                    )
                    images = [Image.open(f) for f in digi_images]
                    resp_text = generate_gemini_content(digi_prompt, active_api_key, working_model_name, images=images)
                    d_blocks = resp_text.split("|||")
                    st.session_state.digi_blocks = [{'id': str(uuid.uuid4()), 'text': b.strip()} for b in d_blocks if b.strip()]
                    st.session_state.digi_saved = False
                    update_paper_count(st.session_state.username)
                    st.rerun()
                except Exception as e:
                    error_msg = str(e).lower()
                    logging.error(f"[Digitize Error] {e}")
                    if "429" in error_msg or "quota" in error_msg:
                        st.error("🚨 The daily generation limit has been reached! Try again later or add your own API key in Advanced Settings.")
                    else:
                        st.error("Couldn't read that paper. Try clearer/well-lit photos, or fewer pages at once.")

    if st.session_state.digi_blocks:
        st.markdown("---")
        st.success(f"Read {len(st.session_state.digi_blocks)} question(s). Review and fix anything the OCR missed below.")
        with st.expander("🛠️ Review & Edit", expanded=True):
            for i, b in enumerate(st.session_state.digi_blocks):
                new_text = st.text_area(f"Question {i+1}", b['text'], height=100, key=f"digi_text_{b['id']}")
                if new_text != st.session_state.digi_blocks[i]['text']:
                    st.session_state.digi_blocks[i]['text'] = new_text
                    st.session_state.digi_saved = False

        digi_md = "\n\n".join([b['text'] for b in st.session_state.digi_blocks])
        digi_html = create_a4_html(digi_md, inst_name, inst_address, inst_contact, teacher_name, inst_logo, is_two_column, digi_subject or "Digitized Paper", digi_class or "N/A", "N/A", exam_time, "")
        digi_word = create_word_docx(digi_md, inst_name, inst_address, inst_contact, teacher_name, inst_logo, is_two_column, digi_subject or "Digitized Paper", digi_class or "N/A", "N/A", exam_time, "")

        gc1, gc2, gc3, gc4 = st.columns(4)
        gc1.download_button("🖨️ HTML", digi_html, f"{digi_subject or 'Digitized'}_Paper.html", "text/html")
        gc2.download_button("📄 Word", digi_word, f"{digi_subject or 'Digitized'}_Paper.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        digi_pdf = html_to_pdf(digi_html)
        if digi_pdf:
            gc3.download_button("📕 PDF", digi_pdf, f"{digi_subject or 'Digitized'}_Paper.pdf", "application/pdf")
        else:
            gc3.caption("PDF unavailable")
        st.caption("💡 If the Word file shows boxes instead of Hindi text, install the free 'Noto Sans Devanagari' font on the computer opening it (one-time setup). Not needed for the PDF.")
        if gc4.button("☁️ Save History", key="digi_save_history"):
            data = {"username": st.session_state.username, "date": datetime.now().strftime("%Y-%m-%d"), "subject": digi_subject or "Digitized Paper", "board": "Digitized", "content": digi_md}
            try:
                supabase.table("papers").insert(data).execute()
                st.session_state.digi_saved = True
                st.success("Saved!")
            except Exception as e:
                logging.error(f"[Digitize Save History Error] {e}")
                st.error("Couldn't save to Cloud History. Please try again.")

with tab_history:
    st.markdown("### Cloud History")
    with st.spinner("Loading your saved papers..."):
        try:
            res = supabase.table("papers").select("*").eq("username", st.session_state.username).order("id", desc=True).execute()
            history_error = None
        except Exception as e:
            logging.error(f"[History Load Error] {e}")
            res = None
            history_error = "Couldn't load your history right now. Please refresh."

    if history_error:
        st.error(history_error)
    elif res.data:
        st.caption(f"{len(res.data)} saved paper(s)")
        for p in res.data:
            with st.expander(f"📄 {p['subject']} ({p['date']})"):
                h_html = create_a4_html(p['content'], inst_name, inst_address, inst_contact, teacher_name, inst_logo, is_two_column, p['subject'], "N/A", "N/A", exam_time, "")
                h_word = create_word_docx(p['content'], inst_name, inst_address, inst_contact, teacher_name, inst_logo, is_two_column, p['subject'], "N/A", "N/A", exam_time, "")

                dl1, dl2, dl3, dl4 = st.columns(4)
                dl1.download_button("🖨️ HTML", h_html, f"History_{p['id']}.html", "text/html", key=f"h_{p['id']}")
                dl2.download_button("📄 Word", h_word, f"History_{p['id']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"w_{p['id']}")
                h_pdf = html_to_pdf(h_html)
                if h_pdf:
                    dl3.download_button("📕 PDF", h_pdf, f"History_{p['id']}.pdf", "application/pdf", key=f"pdf_{p['id']}")
                else:
                    dl3.caption("PDF unavailable")

                confirm_key = f"confirm_del_{p['id']}"
                if confirm_key not in st.session_state: st.session_state[confirm_key] = False

                if not st.session_state[confirm_key]:
                    if dl4.button("🗑️ Delete", key=f"d_{p['id']}"):
                        st.session_state[confirm_key] = True
                        st.rerun()
                else:
                    st.warning("Delete this paper permanently?")
                    yc, nc = st.columns(2)
                    if yc.button("Yes, delete", key=f"yd_{p['id']}"):
                        delete_paper(p['id'], st.session_state.username)
                        st.session_state[confirm_key] = False
                        st.rerun()
                    if nc.button("Cancel", key=f"nd_{p['id']}"):
                        st.session_state[confirm_key] = False
                        st.rerun()
    else:
        st.info("No saved papers yet — generate one and click '☁️ Save History' to keep it here.")
